"""
CleanDoc - Servicio de limpieza de documentos
==============================================
Servicio principal para limpiar documentos DOCX institucionales.

IMPORTANTE: Este servicio SOLO limpia elementos institucionales de los encabezados,
imágenes en headers, textboxes y la sección de firmas. El contenido del documento
permanece INTACTO.
"""

import logging
import re
from io import BytesIO
from typing import Optional, Dict, Any
from dataclasses import dataclass, field

from docx import Document
from lxml import etree

from ..utils.exceptions import FileProcessingError

# Configurar logger
logger = logging.getLogger(__name__)


@dataclass
class CleaningStats:
    """Estadísticas de limpieza de un documento"""
    images_removed: int = 0
    institutional_paragraphs_cleaned: int = 0
    textboxes_cleaned: int = 0
    signature_section_removed: bool = False
    paragraphs_removed: int = 0
    errors: list = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """Convierte las estadísticas a un diccionario"""
        return {
            'images_removed': self.images_removed,
            'institutional_paragraphs_cleaned': self.institutional_paragraphs_cleaned,
            'textboxes_cleaned': self.textboxes_cleaned,
            'signature_section_removed': self.signature_section_removed,
            'paragraphs_removed': self.paragraphs_removed,
            'has_errors': len(self.errors) > 0,
            'error_count': len(self.errors)
        }


class DocumentCleaner:
    """
    Limpiador de documentos DOCX institucionales.

    Este servicio elimina:
    1. Imágenes de encabezados (excepto las que están en tablas)
    2. Texto institucional en párrafos ("ÓRGANO DE FISCALIZACIÓN SUPERIOR", etc.)
    3. Texto institucional en textboxes
    4. Sección de firmas (desde "Elaboró" hasta el final)

    GARANTÍA: El contenido principal del documento permanece INTACTO.
    Solo se modifican encabezados, imágenes de headers, textboxes y firmas.
    """

    # Patrones de búsqueda para texto institucional
    PATTERN_ORGANO = re.compile(
        r"ÓRGANO\s+DE\s+FISCALIZACI[ÓO]N\s+SUPERIOR",
        re.IGNORECASE
    )

    PATTERN_DIRECCION = re.compile(
        r"DIRECCI[ÓO]N\s+DE\s+AUDITOR[IÍ]A\s+A\s+ENTES\s+ESTATALES",
        re.IGNORECASE
    )

    PATTERN_ELABORO = re.compile(
        r"Elabor[oó]",
        re.IGNORECASE
    )

    def __init__(self):
        """Inicializa el limpiador de documentos"""
        self.stats = CleaningStats()

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        """
        Normaliza espacios en blanco en un texto.

        Args:
            text: Texto a normalizar

        Returns:
            Texto con espacios normalizados
        """
        return re.sub(r"\s+", " ", text or "")

    def _remove_header_images(self, doc: Document) -> None:
        """
        Elimina imágenes de los encabezados del documento.

        IMPORTANTE: Preserva imágenes que están dentro de tablas.
        Elimina tanto imágenes DrawingML modernas como imágenes VML antiguas.

        Args:
            doc: Documento a procesar
        """
        logger.debug("Eliminando imágenes de encabezados...")

        for section in doc.sections:
            try:
                # Eliminar imágenes DrawingML modernas (<drawing>)
                drawings = section.header._element.xpath(".//*[local-name()='drawing']")
                for shape in drawings:
                    # Verificar si está dentro de una tabla
                    if shape.xpath("ancestor::*[local-name()='tbl']"):
                        logger.debug("Preservando imagen dentro de tabla")
                        continue

                    parent = shape.getparent()
                    if parent is not None:
                        parent.remove(shape)
                        self.stats.images_removed += 1
                        logger.debug("Imagen DrawingML eliminada")

                # Eliminar imágenes VML antiguas (<pict>)
                # IMPORTANTE: NO eliminar <pict> que contengan textboxes
                picts = section.header._element.xpath(".//*[local-name()='pict']")
                for pict in picts:
                    # Verificar si está dentro de una tabla
                    if pict.xpath("ancestor::*[local-name()='tbl']"):
                        logger.debug("Preservando imagen VML dentro de tabla")
                        continue

                    # NO eliminar si contiene textboxes (contienen texto importante)
                    if pict.xpath(".//*[local-name()='txbxContent']"):
                        logger.debug("Preservando pict con textbox")
                        continue

                    parent = pict.getparent()
                    if parent is not None:
                        parent.remove(pict)
                        self.stats.images_removed += 1
                        logger.debug("Imagen VML eliminada")

            except Exception as e:
                error_msg = f"Error eliminando imágenes de encabezado: {str(e)}"
                logger.error(error_msg)
                self.stats.errors.append(error_msg)

    def _clean_institutional_paragraphs(self, doc: Document) -> None:
        """
        Limpia párrafos que contienen texto institucional.

        Elimina o limpia párrafos que contienen:
        - "ÓRGANO DE FISCALIZACIÓN SUPERIOR"
        - "DIRECCIÓN DE AUDITORÍA A ENTES ESTATALES"

        Args:
            doc: Documento a procesar
        """
        logger.debug("Limpiando párrafos institucionales...")

        for p in list(doc.paragraphs):
            try:
                texto = p.text
                if self.PATTERN_ORGANO.search(texto) or self.PATTERN_DIRECCION.search(texto):
                    # Limpiar el texto completo del párrafo
                    texto_limpio = self.PATTERN_ORGANO.sub("", texto)
                    texto_limpio = self.PATTERN_DIRECCION.sub("", texto_limpio)
                    texto_limpio = texto_limpio.strip()

                    if texto_limpio:
                        # Si queda texto, preservarlo en el primer run
                        if p.runs:
                            p.runs[0].text = texto_limpio
                            for run in p.runs[1:]:
                                run.text = ""
                            self.stats.institutional_paragraphs_cleaned += 1
                            logger.debug(f"Párrafo limpiado: '{texto}' -> '{texto_limpio}'")
                    else:
                        # Si no queda texto, eliminar el párrafo completo
                        p._element.getparent().remove(p._element)
                        self.stats.paragraphs_removed += 1
                        logger.debug(f"Párrafo eliminado: '{texto}'")

            except Exception as e:
                error_msg = f"Error limpiando párrafo: {str(e)}"
                logger.error(error_msg)
                self.stats.errors.append(error_msg)

    def _clean_textboxes(self, xmlroot) -> None:
        """
        Limpia texto institucional dentro de textboxes.

        Args:
            xmlroot: Elemento raíz XML a procesar
        """
        try:
            for p in xmlroot.xpath(".//*[local-name()='txbxContent']//*[local-name()='p']"):
                ts = p.xpath(".//*[local-name()='t']")
                if not ts:
                    continue

                original = " ".join(
                    self._normalize_whitespace(t.text) for t in ts if t.text
                )

                # Aplicar los mismos patrones de limpieza
                nuevo = self.PATTERN_ORGANO.sub("", original)
                nuevo = self.PATTERN_DIRECCION.sub("", nuevo)
                nuevo = nuevo.strip()

                if nuevo != original:
                    ts[0].text = nuevo
                    for t in ts[1:]:
                        t.text = ""

                    self.stats.textboxes_cleaned += 1
                    logger.debug(f"Textbox limpiado: '{original}' -> '{nuevo}'")

                    # Si no queda texto, eliminar el párrafo
                    if not nuevo:
                        parent = p.getparent()
                        if parent is not None:
                            parent.remove(p)

        except Exception as e:
            error_msg = f"Error limpiando textboxes: {str(e)}"
            logger.error(error_msg)
            self.stats.errors.append(error_msg)

    def _remove_signature_section(self, doc: Document) -> None:
        """
        Elimina la sección de firmas del documento.

        Busca el primer "Elaboró" y elimina todo desde ahí hasta el final.

        Args:
            doc: Documento a procesar
        """
        logger.debug("Buscando sección de firmas...")

        try:
            indice_inicio = None
            for i, p in enumerate(doc.paragraphs):
                # Buscar "Elaboró" sin espacios
                if self.PATTERN_ELABORO.search(p.text.replace(" ", "")):
                    indice_inicio = i
                    logger.debug(f"Sección de firmas encontrada en párrafo {i}: '{p.text}'")
                    break

            if indice_inicio is not None:
                # Contar cuántos párrafos se eliminarán
                paragraphs_to_remove = len(doc.paragraphs) - indice_inicio

                # Eliminar todos los párrafos desde "Elaboró" hasta el final
                for j in range(len(doc.paragraphs) - 1, indice_inicio - 1, -1):
                    para = doc.paragraphs[j]
                    para._element.getparent().remove(para._element)

                self.stats.signature_section_removed = True
                self.stats.paragraphs_removed += paragraphs_to_remove
                logger.info(f"Sección de firmas eliminada ({paragraphs_to_remove} párrafos)")
            else:
                logger.debug("No se encontró sección de firmas")

        except Exception as e:
            error_msg = f"Error eliminando sección de firmas: {str(e)}"
            logger.error(error_msg)
            self.stats.errors.append(error_msg)

    def clean_document(
        self,
        file_stream,
        filename: Optional[str] = None
    ) -> tuple[BytesIO, CleaningStats]:
        """
        Limpia un documento DOCX eliminando elementos institucionales.

        Este método:
        1. Elimina imágenes de encabezados (excepto las de tablas)
        2. Limpia texto institucional en párrafos
        3. Limpia texto institucional en textboxes
        4. Elimina la sección de firmas

        GARANTÍA: El contenido principal del documento NO se modifica.

        Args:
            file_stream: Stream del archivo DOCX
            filename: Nombre del archivo (opcional, para logging)

        Returns:
            Tupla con (archivo_limpio, estadísticas)

        Raises:
            FileProcessingError: Si hay un error procesando el documento
        """
        # Reiniciar estadísticas
        self.stats = CleaningStats()

        logger.info(f"Iniciando limpieza de documento: {filename or 'sin nombre'}")

        try:
            # Cargar el documento
            doc = Document(file_stream)
            logger.debug("Documento cargado exitosamente")

            # 1. Eliminar imágenes de encabezados
            self._remove_header_images(doc)

            # 2. Limpiar párrafos institucionales
            self._clean_institutional_paragraphs(doc)

            # 3. Limpiar textboxes en todo el documento
            self._clean_textboxes(doc._element)
            for section in doc.sections:
                self._clean_textboxes(section.header._element)
                self._clean_textboxes(section.footer._element)

            # 4. Eliminar sección de firmas
            self._remove_signature_section(doc)

            # Guardar el documento en memoria
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            logger.info(
                f"Documento limpiado exitosamente: {filename or 'sin nombre'} - "
                f"Estadísticas: {self.stats.to_dict()}"
            )

            return output, self.stats

        except Exception as e:
            error_msg = f"Error procesando documento {filename or 'sin nombre'}: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise FileProcessingError(error_msg, filename=filename)


# Instancia singleton del limpiador
_cleaner_instance: Optional[DocumentCleaner] = None


def get_cleaner() -> DocumentCleaner:
    """
    Obtiene la instancia singleton del limpiador de documentos.

    Returns:
        Instancia de DocumentCleaner
    """
    global _cleaner_instance
    if _cleaner_instance is None:
        _cleaner_instance = DocumentCleaner()
    return _cleaner_instance
