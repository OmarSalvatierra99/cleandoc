from io import BytesIO
from pathlib import Path
import unittest

from docx import Document

from scripts.utils import get_cleaner


LOGO_PATH = Path(__file__).resolve().parents[1] / 'static' / 'img' / 'ofs_logo.png'


class DocumentCleanerTests(unittest.TestCase):
    def _count_header_graphics(self, doc: Document) -> int:
        total = 0
        seen = set()

        for section in doc.sections:
            for header in (
                section.header,
                section.first_page_header,
                section.even_page_header,
            ):
                header_id = id(header._element)
                if header_id in seen:
                    continue
                seen.add(header_id)

                total += len(header._element.xpath(".//*[local-name()='drawing']"))
                total += len(header._element.xpath(".//*[local-name()='pict']"))

        return total

    def _build_doc(self, include_primary_header: bool = True, include_first_page_header: bool = False) -> BytesIO:
        doc = Document()
        doc.add_paragraph('Contenido de prueba')

        section = doc.sections[0]
        if include_primary_header:
            paragraph = section.header.paragraphs[0]
            paragraph.add_run('Texto institucional ')
            paragraph.add_run().add_picture(str(LOGO_PATH))

        if include_first_page_header:
            section.different_first_page_header_footer = True
            first_page_paragraph = section.first_page_header.paragraphs[0]
            first_page_paragraph.add_run().add_picture(str(LOGO_PATH))

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream

    def test_clean_document_removes_header_images_and_preserves_header_text(self):
        cleaner = get_cleaner()
        source = self._build_doc(include_primary_header=True)

        output, stats = cleaner.clean_document(source, 'header.docx')
        cleaned_doc = Document(output)

        self.assertEqual(stats.images_removed, 1)
        self.assertEqual(self._count_header_graphics(cleaned_doc), 0)
        self.assertIn('Texto institucional', cleaned_doc.sections[0].header.paragraphs[0].text)
        self.assertEqual(cleaned_doc.paragraphs[0].text, 'Contenido de prueba')

    def test_clean_document_removes_first_page_header_images(self):
        cleaner = get_cleaner()
        source = self._build_doc(
            include_primary_header=False,
            include_first_page_header=True,
        )

        output, stats = cleaner.clean_document(source, 'first-page.docx')
        cleaned_doc = Document(output)

        self.assertEqual(stats.images_removed, 1)
        self.assertEqual(self._count_header_graphics(cleaned_doc), 0)

    def test_get_cleaner_returns_fresh_instances(self):
        first_cleaner = get_cleaner()
        second_cleaner = get_cleaner()

        self.assertIsNot(first_cleaner, second_cleaner)


if __name__ == '__main__':
    unittest.main()
