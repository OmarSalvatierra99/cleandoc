"""
CleanDoc - Tests para DocumentCleaner
======================================
Tests unitarios para el servicio de limpieza de documentos.
"""

import unittest
from io import BytesIO
from docx import Document

import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from app.services.document_cleaner import DocumentCleaner, CleaningStats


class TestDocumentCleaner(unittest.TestCase):
    """Tests para la clase DocumentCleaner"""

    def setUp(self):
        """Configuración antes de cada test"""
        self.cleaner = DocumentCleaner()

    def test_cleaner_initialization(self):
        """Test: El cleaner se inicializa correctamente"""
        self.assertIsInstance(self.cleaner, DocumentCleaner)
        self.assertIsInstance(self.cleaner.stats, CleaningStats)

    def test_normalize_whitespace(self):
        """Test: La normalización de espacios funciona correctamente"""
        test_cases = [
            ("Hello    World", "Hello World"),
            ("Multiple   \n\t  spaces", "Multiple spaces"),
            ("  Trim  spaces  ", "Trim spaces"),
            ("", ""),
            (None, "")
        ]

        for input_text, expected in test_cases:
            result = self.cleaner._normalize_whitespace(input_text)
            self.assertEqual(result, expected)

    def test_cleaning_stats_to_dict(self):
        """Test: Las estadísticas se convierten correctamente a diccionario"""
        stats = CleaningStats(
            images_removed=5,
            institutional_paragraphs_cleaned=3,
            textboxes_cleaned=2,
            signature_section_removed=True,
            paragraphs_removed=10
        )

        result = stats.to_dict()

        self.assertEqual(result['images_removed'], 5)
        self.assertEqual(result['institutional_paragraphs_cleaned'], 3)
        self.assertEqual(result['textboxes_cleaned'], 2)
        self.assertTrue(result['signature_section_removed'])
        self.assertEqual(result['paragraphs_removed'], 10)
        self.assertFalse(result['has_errors'])
        self.assertEqual(result['error_count'], 0)

    def test_pattern_matching(self):
        """Test: Los patrones regex funcionan correctamente"""
        # Test pattern ORGANO
        test_text_organo = "ÓRGANO DE FISCALIZACIÓN SUPERIOR"
        self.assertTrue(self.cleaner.PATTERN_ORGANO.search(test_text_organo))

        # Test pattern DIRECCION
        test_text_dir = "DIRECCIÓN DE AUDITORÍA A ENTES ESTATALES"
        self.assertTrue(self.cleaner.PATTERN_DIRECCION.search(test_text_dir))

        # Test pattern ELABORO
        test_text_elaboro = "Elaboró:"
        self.assertTrue(self.cleaner.PATTERN_ELABORO.search(test_text_elaboro))

        # Test case insensitive
        test_text_lower = "órgano de fiscalización superior"
        self.assertTrue(self.cleaner.PATTERN_ORGANO.search(test_text_lower))

    def test_clean_simple_document(self):
        """Test: Limpieza de un documento simple"""
        # Crear documento de prueba
        doc = Document()
        doc.add_paragraph("Contenido normal del documento")
        doc.add_paragraph("ÓRGANO DE FISCALIZACIÓN SUPERIOR")
        doc.add_paragraph("Más contenido")
        doc.add_paragraph("Elaboró: Juan Pérez")
        doc.add_paragraph("Este párrafo debe eliminarse")

        # Guardar en memoria
        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        # Limpiar documento
        cleaned_stream, stats = self.cleaner.clean_document(stream)

        # Verificar que se retorna un BytesIO
        self.assertIsInstance(cleaned_stream, BytesIO)

        # Verificar estadísticas
        self.assertIsInstance(stats, CleaningStats)
        self.assertTrue(stats.signature_section_removed)

        # Cargar documento limpio y verificar contenido
        cleaned_stream.seek(0)
        cleaned_doc = Document(cleaned_stream)

        # Verificar que quedan solo los párrafos esperados
        # (El documento limpio debe tener menos párrafos)
        self.assertLess(len(cleaned_doc.paragraphs), len(doc.paragraphs))


class TestCleaningStats(unittest.TestCase):
    """Tests para la clase CleaningStats"""

    def test_stats_default_values(self):
        """Test: Los valores por defecto son correctos"""
        stats = CleaningStats()

        self.assertEqual(stats.images_removed, 0)
        self.assertEqual(stats.institutional_paragraphs_cleaned, 0)
        self.assertEqual(stats.textboxes_cleaned, 0)
        self.assertFalse(stats.signature_section_removed)
        self.assertEqual(stats.paragraphs_removed, 0)
        self.assertEqual(len(stats.errors), 0)

    def test_stats_with_errors(self):
        """Test: Las estadísticas manejan errores correctamente"""
        stats = CleaningStats()
        stats.errors.append("Error de prueba 1")
        stats.errors.append("Error de prueba 2")

        result = stats.to_dict()

        self.assertTrue(result['has_errors'])
        self.assertEqual(result['error_count'], 2)


if __name__ == '__main__':
    unittest.main()
