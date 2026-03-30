"""
CleanDoc - Utilidades y servicios auxiliares
============================================
Funciones de validacion, excepciones y limpieza de documentos.
"""

import logging
import math
import os
import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
from docx import Document
from pypdf import PdfReader, PdfWriter, PageObject, Transformation
from pypdf.generic import ContentStream
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)


class CleanDocError(Exception):
    """Excepcion base para errores de CleanDoc."""

    def __init__(self, message: str, status_code: int = 500):
        self.message = message
        self.status_code = status_code
        super().__init__(self.message)


class InvalidFileError(CleanDocError):
    """Excepcion cuando el archivo no es válido."""

    def __init__(self, message: str = "Archivo inválido"):
        super().__init__(message, status_code=400)


class FileProcessingError(CleanDocError):
    """Excepcion durante el procesamiento del archivo."""

    def __init__(self, message: str = "Error al procesar el archivo", filename: str = None):
        self.filename = filename
        full_message = f"{message}: {filename}" if filename else message
        super().__init__(full_message, status_code=500)


class NoFilesProvidedError(CleanDocError):
    """Excepcion cuando no se proporcionan archivos."""

    def __init__(self, message: str = "No se proporcionaron archivos"):
        super().__init__(message, status_code=400)


class FileTooLargeError(CleanDocError):
    """Excepcion cuando el archivo excede el tamaño máximo."""

    def __init__(self, message: str = "El archivo excede el tamaño máximo permitido"):
        super().__init__(message, status_code=413)


class UnsupportedFileTypeError(CleanDocError):
    """Excepcion cuando el tipo de archivo no esta soportado."""

    def __init__(self, message: str = "Tipo de archivo no soportado. Solo se permiten archivos .docx"):
        super().__init__(message, status_code=415)


def sanitize_filename(filename: str) -> str:
    """Sanitiza el nombre del archivo para prevenir path traversal."""
    if not filename:
        raise InvalidFileError("El nombre del archivo está vacío")

    safe_name = secure_filename(filename)

    if not safe_name:
        raise InvalidFileError(f"Nombre de archivo inválido: {filename}")

    if not re.match(r'^[\w\-. ]+$', safe_name):
        raise InvalidFileError(
            f"Nombre de archivo contiene caracteres no permitidos: {filename}"
        )

    return safe_name


def validate_file_extension(filename: str, allowed_extensions: set = {'.docx'}) -> bool:
    """Valida que el archivo tenga una extensión permitida."""
    if not filename:
        raise InvalidFileError("El nombre del archivo está vacío")

    file_ext = Path(filename).suffix.lower()

    if file_ext not in allowed_extensions:
        raise UnsupportedFileTypeError(
            f"Extensión '{file_ext}' no permitida. Solo se permiten: {', '.join(allowed_extensions)}"
        )

    return True


def validate_pdf_file(file: FileStorage, max_size: int = 50 * 1024 * 1024) -> Tuple[str, bool]:
    """Valida completamente un archivo PDF."""
    if not file or not file.filename:
        raise InvalidFileError("No se proporcionó un archivo PDF válido")

    safe_filename = sanitize_filename(file.filename)
    validate_file_extension(safe_filename, allowed_extensions={'.pdf'})
    validate_file_size(file, max_size)

    return safe_filename, True


def validate_file_size(file: FileStorage, max_size: int = 50 * 1024 * 1024) -> bool:
    """Valida que el archivo no exceda el tamaño máximo."""
    if file is None:
        raise InvalidFileError("El archivo es None")

    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)

    if size > max_size:
        max_size_mb = max_size / (1024 * 1024)
        current_size_mb = size / (1024 * 1024)
        raise FileTooLargeError(
            f"El archivo ({current_size_mb:.2f} MB) excede el tamaño máximo permitido ({max_size_mb:.2f} MB)"
        )

    return True


def validate_docx_file(file: FileStorage, max_size: int = 50 * 1024 * 1024) -> tuple[str, bool]:
    """Valida completamente un archivo DOCX."""
    if not file or not file.filename:
        raise InvalidFileError("No se proporcionó un archivo válido")

    safe_filename = sanitize_filename(file.filename)
    validate_file_extension(safe_filename)
    validate_file_size(file, max_size)

    return safe_filename, True


def is_valid_docx_content(file_stream) -> bool:
    """Valida que el contenido sea un DOCX verificando firma ZIP."""
    try:
        file_stream.seek(0)
        header = file_stream.read(4)
        file_stream.seek(0)

        zip_signatures = [
            b'PK\x03\x04',
            b'PK\x05\x06',
            b'PK\x07\x08',
        ]

        return any(header.startswith(sig) for sig in zip_signatures)

    except Exception:
        return False


def is_valid_pdf_content(file_stream) -> bool:
    """Valida que el contenido sea un PDF verificando su encabezado."""
    try:
        file_stream.seek(0)
        header = file_stream.read(5)
        file_stream.seek(0)
        return header.startswith(b'%PDF-')
    except Exception:
        return False


@dataclass
class PdfScalingStats:
    """Estadísticas de escalado de un PDF."""
    total_pages: int = 0
    adjusted_pages: int = 0
    upscaled_pages: int = 0
    output_pages: int = 0
    avg_scale: float = 1.0
    min_scale: float = 1.0
    max_scale: float = 1.0
    target_size: str = "LETTER"
    margin_ratio: float = 0.92
    n_up: int = 1

    def to_dict(self) -> Dict[str, Any]:
        return {
            'total_pages': self.total_pages,
            'adjusted_pages': self.adjusted_pages,
            'upscaled_pages': self.upscaled_pages,
            'output_pages': self.output_pages,
            'avg_scale': round(self.avg_scale, 4),
            'min_scale': round(self.min_scale, 4),
            'max_scale': round(self.max_scale, 4),
            'target_size': self.target_size,
            'margin_ratio': round(self.margin_ratio, 4),
            'n_up': self.n_up,
        }


PDF_PAGE_SIZES = {
    "LETTER": (612.0, 792.0),  # 8.5 x 11 in
    "A4": (595.28, 841.89),    # 210 x 297 mm
}


def _page_has_table_lines(page: PageObject, reader: PdfReader) -> bool:
    """Detecta líneas de tablas buscando operadores de trazo en el contenido."""
    try:
        contents = page.get_contents()
        if not contents:
            return False
        content_stream = ContentStream(contents, reader)
        line_ops = 0
        rect_ops = 0

        for operands, operator in content_stream.operations:
            if operator == b"re":
                rect_ops += 1
            elif operator in (b"l", b"m", b"S", b"s", b"B", b"b", b"B*", b"b*"):
                line_ops += 1

        return rect_ops >= 3 or line_ops >= 30
    except Exception:
        return False


def scale_pdf_for_print(
    file_stream,
    target_size: str = "LETTER",
    allow_upscale: bool = False,
    margin_ratio: float = 0.92,
    n_up: int = 1,
) -> tuple[BytesIO, PdfScalingStats]:
    """Escala un PDF para impresión sin recortar contenido."""
    if target_size not in PDF_PAGE_SIZES:
        raise InvalidFileError("Tamaño de página no soportado")
    if margin_ratio <= 0 or margin_ratio > 1:
        raise InvalidFileError("El margen de impresión no es válido")

    try:
        reader = PdfReader(file_stream)
        if reader.is_encrypted:
            if not reader.decrypt(""):
                raise InvalidFileError("El PDF está protegido con contraseña")

        stats = PdfScalingStats()
        stats.target_size = target_size
        stats.margin_ratio = margin_ratio
        stats.n_up = n_up
        total_scale = 0.0
        min_scale = None
        max_scale = None

        writer = PdfWriter()

        slots_per_page = n_up
        slot_cols = 1 if n_up == 2 else 2 if n_up == 4 else 1
        slot_rows = 1 if n_up == 1 else 2
        slot_width = PDF_PAGE_SIZES[target_size][0] / slot_cols
        slot_height = PDF_PAGE_SIZES[target_size][1] / slot_rows
        slot_index = 0
        output_page = None

        for page in reader.pages:
            stats.total_pages += 1
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            has_table_lines = _page_has_table_lines(page, reader)

            page_target_width, page_target_height = PDF_PAGE_SIZES[target_size]
            if n_up == 1 and width > height:
                page_target_width, page_target_height = page_target_height, page_target_width
                slot_width = page_target_width
                slot_height = page_target_height
            elif n_up == 1:
                slot_width = page_target_width
                slot_height = page_target_height

            content_width = slot_width * margin_ratio
            content_height = slot_height * margin_ratio

            scale_uniform = min(content_width / width, content_height / height)
            if not allow_upscale:
                scale_uniform = min(scale_uniform, 1.0)
            scale_x = scale_uniform
            scale_y = scale_uniform

            if output_page is None:
                output_page = PageObject.create_blank_page(
                    width=page_target_width,
                    height=page_target_height,
                )

            slot_col = slot_index % slot_cols
            slot_row = slot_index // slot_cols
            slot_x = slot_col * slot_width
            slot_y = page_target_height - (slot_row + 1) * slot_height

            tx = slot_x + (slot_width - width * scale_x) / 2
            ty = slot_y + (slot_height - height * scale_y) / 2
            output_page.merge_transformed_page(
                page,
                Transformation().scale(scale_x, scale_y).translate(tx, ty),
            )

            slot_index += 1
            if slot_index >= slots_per_page:
                writer.add_page(output_page)
                output_page = None
                slot_index = 0

            if (
                width != slot_width
                or height != slot_height
                or scale_x != 1.0
                or scale_y != 1.0
            ):
                stats.adjusted_pages += 1
            if scale_x > 1.0 or scale_y > 1.0:
                stats.upscaled_pages += 1

            scale_for_stats = min(scale_x, scale_y)
            total_scale += scale_for_stats
            min_scale = scale_for_stats if min_scale is None else min(min_scale, scale_for_stats)
            max_scale = scale_for_stats if max_scale is None else max(max_scale, scale_for_stats)

        if output_page is not None:
            writer.add_page(output_page)

        if stats.total_pages > 0:
            stats.avg_scale = total_scale / stats.total_pages
            stats.min_scale = min_scale if min_scale is not None else 1.0
            stats.max_scale = max_scale if max_scale is not None else 1.0
            stats.output_pages = len(writer.pages)

        output = BytesIO()
        writer.write(output)
        output.seek(0)
        return output, stats

    except CleanDocError:
        raise

    except Exception as e:
        error_msg = f"Error escalando PDF: {str(e)}"
        logger.error(error_msg, exc_info=True)
        raise FileProcessingError(error_msg)


@dataclass
class CleaningStats:
    """Estadísticas de limpieza de un documento."""
    images_removed: int = 0
    institutional_paragraphs_cleaned: int = 0
    textboxes_cleaned: int = 0
    signature_section_removed: bool = False
    paragraphs_removed: int = 0
    errors: list = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            'images_removed': self.images_removed,
            'institutional_paragraphs_cleaned': self.institutional_paragraphs_cleaned,
            'textboxes_cleaned': self.textboxes_cleaned,
            'signature_section_removed': self.signature_section_removed,
            'paragraphs_removed': self.paragraphs_removed,
            'has_errors': len(self.errors) > 0,
            'error_count': len(self.errors)
        }


class DocumentCleaner:
    """Limpiador de documentos DOCX institucionales."""

    PATTERN_ORGANO = re.compile(
        r"ÓRGANO\s+DE\s+FISCALIZACI[ÓO]N\s+SUPERIOR",
        re.IGNORECASE
    )

    PATTERN_DIRECCION = re.compile(
        r"DIRECCI[ÓO]N\s+DE\s+AUDITOR[IÍ]A\s+A\s+ENTES\s+ESTATALES",
        re.IGNORECASE
    )

    PATTERN_ELABORO = re.compile(
        r"Elabor[oó]",
        re.IGNORECASE
    )

    def __init__(self):
        self.stats = CleaningStats()

    @staticmethod
    def _normalize_whitespace(text: str) -> str:
        return re.sub(r"\s+", " ", text or "")

    @staticmethod
    def _iter_document_parts(doc: Document, prefix: str) -> list:
        parts = []
        seen = set()

        # Recorremos las partes OPC del paquete para evitar la recursión
        # infinita que puede disparar python-docx al resolver headers/footers
        # linked-to-previous en documentos complejos.
        for part in doc.part.package.parts:
            partname = str(getattr(part, 'partname', ''))
            if not partname.startswith(prefix) or not partname.endswith('.xml'):
                continue

            element = getattr(part, '_element', None)
            if element is None:
                continue

            element_id = id(element)
            if element_id in seen:
                continue
            seen.add(element_id)
            parts.append(element)

        return parts

    @classmethod
    def _iter_section_headers(cls, doc: Document) -> list:
        return cls._iter_document_parts(doc, '/word/header')

    @classmethod
    def _iter_section_footers(cls, doc: Document) -> list:
        return cls._iter_document_parts(doc, '/word/footer')

    @staticmethod
    def _run_has_non_graphic_content(run_element) -> bool:
        visible_markers = (
            ".//*[local-name()='t' or local-name()='tab' or local-name()='br' "
            "or local-name()='cr' or local-name()='instrText' "
            "or local-name()='delText' or local-name()='sym']"
        )

        for node in run_element.xpath(visible_markers):
            text = getattr(node, 'text', None)
            if text is None or text.strip():
                return True

        return False

    def _resolve_graphic_removal_target(self, graphic_element):
        if graphic_element.xpath("ancestor::*[local-name()='tbl']"):
            return None

        if (
            graphic_element.xpath("ancestor::*[local-name()='txbxContent']")
            or graphic_element.xpath(".//*[local-name()='txbxContent']")
        ):
            return None

        run_ancestors = graphic_element.xpath("ancestor::*[local-name()='r'][1]")
        if not run_ancestors:
            return graphic_element

        run_element = run_ancestors[0]
        if self._run_has_non_graphic_content(run_element):
            return graphic_element

        paragraph_ancestors = run_element.xpath("ancestor::*[local-name()='p'][1]")
        if not paragraph_ancestors:
            return run_element

        paragraph_element = paragraph_ancestors[0]
        paragraph_has_text = bool(
            paragraph_element.xpath(
                ".//*[local-name()='t' or local-name()='instrText' "
                "or local-name()='delText' or local-name()='sym']"
            )
        )

        if paragraph_has_text:
            return run_element

        return paragraph_element

    def _remove_graphics_from_header(self, header_element) -> None:
        removable_graphics = (
            header_element.xpath(".//*[local-name()='drawing']")
            + header_element.xpath(".//*[local-name()='pict']")
        )

        targets = {}

        for graphic in removable_graphics:
            try:
                target = self._resolve_graphic_removal_target(graphic)
                if target is None:
                    continue

                key = id(target)
                if key not in targets:
                    targets[key] = {
                        'element': target,
                        'count': 0,
                    }
                targets[key]['count'] += 1

            except Exception as e:
                error_msg = f"Error evaluando imagen de encabezado: {str(e)}"
                logger.error(error_msg)
                self.stats.errors.append(error_msg)

        for item in targets.values():
            target = item['element']
            parent = target.getparent()
            if parent is None:
                continue

            try:
                parent.remove(target)
                self.stats.images_removed += item['count']
            except Exception as e:
                error_msg = f"Error eliminando imagen de encabezado: {str(e)}"
                logger.error(error_msg)
                self.stats.errors.append(error_msg)

    def _remove_header_images(self, doc: Document) -> None:
        for header in self._iter_section_headers(doc):
            try:
                self._remove_graphics_from_header(header)
            except Exception as e:
                error_msg = f"Error eliminando imágenes de encabezado: {str(e)}"
                logger.error(error_msg)
                self.stats.errors.append(error_msg)

    def _clean_institutional_paragraphs(self, doc: Document) -> None:
        for p in list(doc.paragraphs):
            try:
                texto = p.text
                if self.PATTERN_ORGANO.search(texto) or self.PATTERN_DIRECCION.search(texto):
                    texto_limpio = self.PATTERN_ORGANO.sub("", texto)
                    texto_limpio = self.PATTERN_DIRECCION.sub("", texto_limpio).strip()

                    if texto_limpio:
                        if p.runs:
                            p.runs[0].text = texto_limpio
                            for run in p.runs[1:]:
                                run.text = ""
                            self.stats.institutional_paragraphs_cleaned += 1
                    else:
                        p._element.getparent().remove(p._element)
                        self.stats.paragraphs_removed += 1

            except Exception as e:
                error_msg = f"Error limpiando párrafo: {str(e)}"
                logger.error(error_msg)
                self.stats.errors.append(error_msg)

    def _clean_textboxes(self, xmlroot) -> None:
        try:
            for p in xmlroot.xpath(".//*[local-name()='txbxContent']//*[local-name()='p']"):
                ts = p.xpath(".//*[local-name()='t']")
                if not ts:
                    continue

                original = " ".join(
                    self._normalize_whitespace(t.text) for t in ts if t.text
                )

                nuevo = self.PATTERN_ORGANO.sub("", original)
                nuevo = self.PATTERN_DIRECCION.sub("", nuevo).strip()

                if nuevo != original:
                    ts[0].text = nuevo
                    for t in ts[1:]:
                        t.text = ""

                    self.stats.textboxes_cleaned += 1

                    if not nuevo:
                        parent = p.getparent()
                        if parent is not None:
                            parent.remove(p)

        except Exception as e:
            error_msg = f"Error limpiando textboxes: {str(e)}"
            logger.error(error_msg)
            self.stats.errors.append(error_msg)

    def _remove_signature_section(self, doc: Document) -> None:
        try:
            indice_inicio = None
            paragraphs = list(doc.paragraphs)

            for i, p in enumerate(paragraphs):
                if self.PATTERN_ELABORO.search(p.text.replace(" ", "")):
                    indice_inicio = i
                    break

            if indice_inicio is not None:
                paragraphs_to_remove = len(paragraphs) - indice_inicio
                for para in reversed(paragraphs[indice_inicio:]):
                    para._element.getparent().remove(para._element)

                self.stats.signature_section_removed = True
                self.stats.paragraphs_removed += paragraphs_to_remove

        except Exception as e:
            error_msg = f"Error eliminando sección de firmas: {str(e)}"
            logger.error(error_msg)
            self.stats.errors.append(error_msg)

    def clean_document(self, file_stream, filename: Optional[str] = None) -> tuple[BytesIO, CleaningStats]:
        self.stats = CleaningStats()

        try:
            doc = Document(file_stream)
            self._remove_header_images(doc)
            self._clean_institutional_paragraphs(doc)
            self._clean_textboxes(doc._element)
            for header in self._iter_section_headers(doc):
                self._clean_textboxes(header)
            for footer in self._iter_section_footers(doc):
                self._clean_textboxes(footer)
            self._remove_signature_section(doc)

            output = BytesIO()
            doc.save(output)
            output.seek(0)

            return output, self.stats

        except Exception as e:
            error_msg = f"Error procesando documento {filename or 'sin nombre'}: {str(e)}"
            logger.error(error_msg, exc_info=True)
            raise FileProcessingError(error_msg, filename=filename)

def get_cleaner() -> DocumentCleaner:
    """Crea una instancia fresca del limpiador para cada procesamiento."""
    return DocumentCleaner()
