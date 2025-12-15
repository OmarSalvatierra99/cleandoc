import os
import re
import logging
import tempfile
import zipfile
from io import BytesIO
from pathlib import Path
from typing import List, Tuple
from dataclasses import dataclass, field
from logging.handlers import RotatingFileHandler

from flask import Flask, render_template_string, request, send_file, jsonify
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from werkzeug.exceptions import HTTPException
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

log_dir = Path('logs')
log_dir.mkdir(exist_ok=True)
handler = RotatingFileHandler('logs/cleandoc.log', maxBytes=10*1024*1024, backupCount=10)
handler.setFormatter(logging.Formatter('[%(asctime)s] %(levelname)s: %(message)s'))
app.logger.addHandler(handler)
app.logger.setLevel(logging.INFO)

@dataclass
class CleaningStats:
    images_removed: int = 0
    institutional_paragraphs_cleaned: int = 0
    textboxes_cleaned: int = 0
    signature_section_removed: bool = False
    paragraphs_removed: int = 0
    errors: list = field(default_factory=list)

class DocumentCleaner:
    PATTERN_ORGANO = re.compile(r"ÓRGANO\s+DE\s+FISCALIZACI[ÓO]N\s+SUPERIOR", re.IGNORECASE)
    PATTERN_DIRECCION = re.compile(r"DIRECCI[ÓO]N\s+DE\s+AUDITOR[IÍ]A\s+A\s+ENTES\s+ESTATALES", re.IGNORECASE)
    PATTERN_ELABORO = re.compile(r"Elabor[oó]", re.IGNORECASE)

    def __init__(self):
        self.stats = CleaningStats()

    def clean_document(self, file_stream, filename=None):
        self.stats = CleaningStats()
        try:
            doc = Document(file_stream)
            self._remove_header_images(doc)
            self._clean_institutional_paragraphs(doc)
            self._clean_textboxes(doc._element)
            for section in doc.sections:
                self._clean_textboxes(section.header._element)
                self._clean_textboxes(section.footer._element)
            self._remove_signature_section(doc)
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output, self.stats
        except Exception as e:
            app.logger.error(f"Error processing {filename}: {e}")
            raise

    def _remove_header_images(self, doc):
        for section in doc.sections:
            try:
                drawings = section.header._element.xpath(".//*[local-name()='drawing']")
                for shape in drawings:
                    if not shape.xpath("ancestor::*[local-name()='tbl']"):
                        parent = shape.getparent()
                        if parent is not None:
                            parent.remove(shape)
                            self.stats.images_removed += 1
                picts = section.header._element.xpath(".//*[local-name()='pict']")
                for pict in picts:
                    if not pict.xpath("ancestor::*[local-name()='tbl']") and not pict.xpath(".//*[local-name()='txbxContent']"):
                        parent = pict.getparent()
                        if parent is not None:
                            parent.remove(pict)
                            self.stats.images_removed += 1
            except Exception as e:
                self.stats.errors.append(str(e))

    def _clean_institutional_paragraphs(self, doc):
        for p in list(doc.paragraphs):
            try:
                texto = p.text
                if self.PATTERN_ORGANO.search(texto) or self.PATTERN_DIRECCION.search(texto):
                    texto_limpio = self.PATTERN_ORGANO.sub("", texto)
                    texto_limpio = self.PATTERN_DIRECCION.sub("", texto_limpio).strip()
                    if texto_limpio:
                        if p.runs:
                            p.runs[0].text = texto_limpio
                            for run in p.runs[1:]:
                                run.text = ""
                        self.stats.institutional_paragraphs_cleaned += 1
                    else:
                        p._element.getparent().remove(p._element)
                        self.stats.paragraphs_removed += 1
            except Exception as e:
                self.stats.errors.append(str(e))

    def _clean_textboxes(self, xmlroot):
        try:
            for p in xmlroot.xpath(".//*[local-name()='txbxContent']//*[local-name()='p']"):
                ts = p.xpath(".//*[local-name()='t']")
                if not ts:
                    continue
                original = " ".join(re.sub(r"\s+", " ", t.text or "") for t in ts)
                nuevo = self.PATTERN_ORGANO.sub("", original)
                nuevo = self.PATTERN_DIRECCION.sub("", nuevo).strip()
                if nuevo != original:
                    ts[0].text = nuevo
                    for t in ts[1:]:
                        t.text = ""
                    self.stats.textboxes_cleaned += 1
                    if not nuevo:
                        parent = p.getparent()
                        if parent is not None:
                            parent.remove(p)
        except Exception as e:
            self.stats.errors.append(str(e))

    def _remove_signature_section(self, doc):
        try:
            indice_inicio = None
            for i, p in enumerate(doc.paragraphs):
                if self.PATTERN_ELABORO.search(p.text.replace(" ", "")):
                    indice_inicio = i
                    break
            if indice_inicio is not None:
                paragraphs_to_remove = len(doc.paragraphs) - indice_inicio
                for j in range(len(doc.paragraphs) - 1, indice_inicio - 1, -1):
                    para = doc.paragraphs[j]
                    para._element.getparent().remove(para._element)
                self.stats.signature_section_removed = True
                self.stats.paragraphs_removed += paragraphs_to_remove
        except Exception as e:
            self.stats.errors.append(str(e))

HTML_TEMPLATE = '''<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>CleanDoc · OFS Tlaxcala</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Inter',sans-serif;background:#fafafa;color:#1e293b;min-height:100vh;display:flex;flex-direction:column;line-height:1.6}
.hidden{display:none!important}
code{background:#f8fafc;padding:.2rem .4rem;border-radius:4px;font-size:.9em;color:#3b82f6}
.header-fixed{background:rgba(255,255,255,.95);border-bottom:1px solid #e2e8f0;box-shadow:0 1px 3px rgba(0,0,0,.05);padding:1.5rem 2rem;display:flex;align-items:center;justify-content:center;position:sticky;top:0;z-index:100;backdrop-filter:blur(10px)}
.header-content{display:flex;align-items:center;gap:1.25rem;max-width:1200px;width:100%}
.ofs-logo{height:60px;width:auto;object-fit:contain}
.header-text h1{color:#1e293b;font-weight:700;font-size:1.75rem;letter-spacing:-.5px}
.header-text h1 .version{font-size:.65rem;font-weight:600;color:#3b82f6;background:#eff6ff;padding:.25rem .5rem;border-radius:6px;margin-left:.5rem;vertical-align:middle}
.header-text p{color:#64748b;font-size:.875rem;margin-top:.25rem}
.main-content{flex:1;display:flex;align-items:center;justify-content:center;padding:4rem 1.5rem}
.tool-card{background:#fff;border-radius:20px;box-shadow:0 4px 20px rgba(0,0,0,.08);border:1px solid #e2e8f0;padding:3.5rem;max-width:700px;width:100%;transition:all .3s cubic-bezier(.4,0,.2,1)}
.tool-card:hover{box-shadow:0 10px 40px rgba(0,0,0,.12);transform:translateY(-2px)}
.tool-card h2{color:#1e293b;font-weight:700;font-size:1.875rem;margin-bottom:1rem;letter-spacing:-.5px;text-align:center}
.tool-card h3{color:#1e293b;font-weight:600;font-size:1rem;margin:1.5rem 0 1rem}
.upload-area{background:#f8fafc;border:2px dashed #e2e8f0;border-radius:12px;padding:3rem 2rem;text-align:center;cursor:pointer;transition:all .3s;margin-top:2rem}
.upload-area:hover{border-color:#3b82f6;background:#f0f7ff}
.upload-area.drag-over{border-color:#3b82f6;background:#dbeafe;transform:scale(1.02)}
.upload-icon{font-size:3rem;margin-bottom:1rem;animation:float 3s ease-in-out infinite}
@keyframes float{0%,100%{transform:translateY(0)}50%{transform:translateY(-10px)}}
.upload-area h3{color:#1e293b;font-weight:600;font-size:1.25rem;margin:0 0 .5rem}
.upload-area p{color:#64748b;font-size:.95rem;margin-bottom:1rem}
.file-types{display:inline-block;font-size:.8rem;color:#64748b;background:#fff;padding:.4rem .8rem;border-radius:6px;border:1px solid #e2e8f0}
.file-preview{margin-top:2rem;animation:fadeIn .3s}
@keyframes fadeIn{from{opacity:0;transform:translateY(-10px)}to{opacity:1;transform:translateY(0)}}
.file-list{background:#f8fafc;border-radius:12px;padding:1rem;max-height:300px;overflow-y:auto}
.file-item{display:flex;align-items:center;justify-content:space-between;background:#fff;border:1px solid #e2e8f0;border-radius:12px;padding:.75rem 1rem;margin-bottom:.75rem;transition:all .15s}
.file-item:last-child{margin-bottom:0}
.file-item:hover{box-shadow:0 1px 3px rgba(0,0,0,.05);transform:translateX(4px)}
.file-info{display:flex;align-items:center;gap:.75rem;flex:1;min-width:0}
.file-icon{font-size:1.5rem;flex-shrink:0}
.file-details{display:flex;flex-direction:column;min-width:0;flex:1}
.file-name{font-weight:500;color:#1e293b;font-size:.9rem;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.file-size{font-size:.8rem;color:#64748b}
.btn-remove{background:#ef4444;color:#fff;border:none;border-radius:50%;width:28px;height:28px;display:flex;align-items:center;justify-content:center;cursor:pointer;font-size:1rem;padding:0;transition:all .15s;flex-shrink:0}
.btn-remove:hover{background:#dc2626;transform:scale(1.1)}
.btn-primary{width:100%;background:#3b82f6;color:#fff;border:none;border-radius:12px;padding:1rem 2rem;font-weight:600;font-size:1rem;cursor:pointer;transition:all .3s;margin-top:1rem;box-shadow:0 1px 3px rgba(0,0,0,.05);display:flex;align-items:center;justify-content:center;gap:.5rem}
.btn-primary:hover{background:#2563eb;transform:translateY(-1px);box-shadow:0 4px 20px rgba(0,0,0,.08)}
.btn-primary:active{transform:translateY(0)}
.btn-primary:disabled{opacity:.6;cursor:not-allowed;transform:none}
.btn-primary:disabled:hover{background:#3b82f6;box-shadow:0 1px 3px rgba(0,0,0,.05)}
.btn-icon{font-size:1.2rem}
.progress-container{margin-top:2rem;animation:fadeIn .3s}
.progress-header{display:flex;justify-content:space-between;align-items:center;margin-bottom:.75rem}
#progressText{font-size:.9rem;font-weight:500;color:#1e293b}
#progressPercent{font-size:.9rem;font-weight:600;color:#3b82f6}
.progress-bar{height:8px;background:#f8fafc;border-radius:10px;overflow:hidden;position:relative}
.progress-fill{height:100%;background:linear-gradient(90deg,#3b82f6,#60a5fa);border-radius:10px;transition:width .3s;position:relative;overflow:hidden}
.progress-fill::after{content:'';position:absolute;top:0;left:0;bottom:0;right:0;background:linear-gradient(90deg,transparent,rgba(255,255,255,.3),transparent);animation:shimmer 1.5s infinite}
@keyframes shimmer{0%{transform:translateX(-100%)}100%{transform:translateX(100%)}}
.status-message{margin-top:1.5rem;padding:1rem 1.25rem;border-radius:12px;font-size:.9rem;font-weight:500;text-align:center;animation:fadeIn .3s}
.status-message.status-success{background:#d1fae5;color:#065f46;border:1px solid #6ee7b7}
.status-message.status-error{background:#fee2e2;color:#991b1b;border:1px solid #fca5a5}
.stats-container{margin-top:2rem;padding:1.5rem;background:linear-gradient(135deg,#667eea 0%,#764ba2 100%);border-radius:20px;animation:fadeIn .5s}
.stats-container h3{color:#fff;text-align:center;margin:0 0 1.5rem;font-size:1.1rem}
.stats-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(120px,1fr));gap:1rem}
.stat-item{background:rgba(255,255,255,.95);border-radius:12px;padding:1.25rem 1rem;text-align:center;transition:all .15s}
.stat-item:hover{transform:translateY(-4px);box-shadow:0 4px 20px rgba(0,0,0,.08)}
.stat-value{font-size:2rem;font-weight:800;color:#1e293b;line-height:1;margin-bottom:.5rem}
.stat-label{font-size:.8rem;color:#64748b;font-weight:500;line-height:1.3}
.dashboard-footer{text-align:center;background:#fff;border-top:1px solid #e2e8f0;color:#64748b;font-size:.8rem;padding:1.5rem}
@media(max-width:768px){.main-content{padding:2.5rem 1rem}.tool-card{padding:2.5rem 2rem}.stats-grid{grid-template-columns:repeat(2,1fr)}}
</style>
</head>
<body>
<header class="header-fixed">
<div class="header-content">
<img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg==" alt="OFS" class="ofs-logo">
<div class="header-text">
<h1>CleanDoc <span class="version">v2.0</span></h1>
<p>Sistema de limpieza institucional de documentos Word (.docx)</p>
</div>
</div>
</header>
<main class="main-content">
<section class="tool-card">
<h2>Limpieza automatizada de cédulas</h2>
<p><strong>CleanDoc</strong> depura archivos Word eliminando encabezados institucionales, imágenes y secciones de firmas.</p>
<div class="upload-area" id="uploadArea">
<div class="upload-icon">📁</div>
<h3>Arrastra archivos aquí</h3>
<p>o haz clic para seleccionar</p>
<input type="file" id="fileInput" name="archivo" accept=".docx" multiple hidden>
<span class="file-types">Solo .docx (máx. 50 MB)</span>
</div>
<div id="filePreview" class="file-preview hidden">
<h3>Archivos seleccionados (<span id="fileCount">0</span>)</h3>
<div id="fileList" class="file-list"></div>
<button type="button" id="processBtn" class="btn-primary"><span class="btn-icon">🔧</span>Procesar y limpiar</button>
</div>
<div id="progressContainer" class="progress-container hidden">
<div class="progress-header">
<span id="progressText">Procesando...</span>
<span id="progressPercent">0%</span>
</div>
<div class="progress-bar"><div id="progressFill" class="progress-fill"></div></div>
</div>
<div id="statusMessage" class="status-message hidden"></div>
<div id="statsContainer" class="stats-container hidden">
<h3>Estadísticas de limpieza</h3>
<div class="stats-grid" id="statsGrid"></div>
</div>
</section>
</main>
<footer class="dashboard-footer">© Órgano de Fiscalización Superior del Estado de Tlaxcala — CleanDoc v2.0</footer>
<script>
const uploadArea=document.getElementById("uploadArea"),fileInput=document.getElementById("fileInput"),filePreview=document.getElementById("filePreview"),fileList=document.getElementById("fileList"),fileCount=document.getElementById("fileCount"),processBtn=document.getElementById("processBtn"),progressContainer=document.getElementById("progressContainer"),progressFill=document.getElementById("progressFill"),progressText=document.getElementById("progressText"),progressPercent=document.getElementById("progressPercent"),statusMessage=document.getElementById("statusMessage"),statsContainer=document.getElementById("statsContainer"),statsGrid=document.getElementById("statsGrid");let selectedFiles=[];uploadArea.addEventListener("click",()=>fileInput.click());uploadArea.addEventListener("dragover",e=>{e.preventDefault();uploadArea.classList.add("drag-over")});uploadArea.addEventListener("dragleave",()=>uploadArea.classList.remove("drag-over"));uploadArea.addEventListener("drop",e=>{e.preventDefault();uploadArea.classList.remove("drag-over");const files=Array.from(e.dataTransfer.files).filter(file=>file.name.toLowerCase().endsWith('.docx'));files.length>0?handleFiles(files):showStatus("Solo se permiten archivos .docx","error")});fileInput.addEventListener("change",e=>handleFiles(Array.from(e.target.files)));function handleFiles(files){selectedFiles=files;displayFilePreview();hideStatus();hideStats()}function displayFilePreview(){if(selectedFiles.length===0){filePreview.classList.add("hidden");return}filePreview.classList.remove("hidden");fileCount.textContent=selectedFiles.length;fileList.innerHTML="";selectedFiles.forEach((file,index)=>fileList.appendChild(createFileItem(file,index)))}function createFileItem(file,index){const item=document.createElement("div");item.className="file-item";const fileSize=formatFileSize(file.size);item.innerHTML=`<div class="file-info"><span class="file-icon">📄</span><div class="file-details"><span class="file-name">${escapeHtml(file.name)}</span><span class="file-size">${fileSize}</span></div></div><button type="button" class="btn-remove" data-index="${index}">✕</button>`;item.querySelector(".btn-remove").addEventListener("click",()=>removeFile(index));return item}function removeFile(index){selectedFiles.splice(index,1);displayFilePreview();if(selectedFiles.length===0)fileInput.value=""}function formatFileSize(bytes){if(bytes===0)return"0 Bytes";const k=1024,sizes=["Bytes","KB","MB","GB"],i=Math.floor(Math.log(bytes)/Math.log(k));return Math.round(bytes/Math.pow(k,i)*100)/100+" "+sizes[i]}function escapeHtml(text){const div=document.createElement('div');div.textContent=text;return div.innerHTML}processBtn.addEventListener("click",async()=>{if(selectedFiles.length===0){showStatus("Selecciona al menos un archivo .docx","error");return}processBtn.disabled=true;showProgress();hideStatus();hideStats();const formData=new FormData();selectedFiles.forEach(file=>formData.append("archivo",file));try{updateProgress(10,"Enviando...");const response=await fetch("/limpiar_cedula",{method:"POST",body:formData});updateProgress(50,"Procesando...");if(!response.ok){const data=await response.json().catch(()=>({}));throw new Error(data.error||"Error al procesar")}updateProgress(80,"Preparando descarga...");const stats=extractStats(response.headers),blob=await response.blob(),url=URL.createObjectURL(blob),a=document.createElement("a");a.href=url;const contentDisposition=response.headers.get("Content-Disposition");a.download=contentDisposition?contentDisposition.split("filename=")[1]?.replace(/"/g,""):"resultado.zip";document.body.appendChild(a);a.click();document.body.removeChild(a);URL.revokeObjectURL(url);updateProgress(100,"¡Completado!");setTimeout(()=>{hideProgress();showStatus("✅ Descarga completada","success");displayStats(stats);resetForm()},500)}catch(error){hideProgress();showStatus(`❌ ${error.message}`,"error");processBtn.disabled=false}});function showProgress(){progressContainer.classList.remove("hidden")}function hideProgress(){progressContainer.classList.add("hidden")}function updateProgress(percent,text){progressFill.style.width=`${percent}%`;progressPercent.textContent=`${percent}%`;progressText.textContent=text}function showStatus(message,type){statusMessage.textContent=message;statusMessage.className=`status-message status-${type}`;statusMessage.classList.remove("hidden")}function hideStatus(){statusMessage.classList.add("hidden")}function extractStats(headers){return{imagesRemoved:parseInt(headers.get("X-CleanDoc-Images-Removed")||0),paragraphsCleaned:parseInt(headers.get("X-CleanDoc-Paragraphs-Cleaned")||0),signatureRemoved:headers.get("X-CleanDoc-Signature-Removed")==="True",totalFiles:parseInt(headers.get("X-CleanDoc-Total-Files")||1),totalImages:parseInt(headers.get("X-CleanDoc-Total-Images-Removed")||0),totalParagraphs:parseInt(headers.get("X-CleanDoc-Total-Paragraphs-Cleaned")||0)}}function displayStats(stats){const hasMultipleFiles=stats.totalFiles>1;statsGrid.innerHTML=`${hasMultipleFiles?`<div class="stat-item"><div class="stat-value">${stats.totalFiles}</div><div class="stat-label">Archivos procesados</div></div>`:''}<div class="stat-item"><div class="stat-value">${hasMultipleFiles?stats.totalImages:stats.imagesRemoved}</div><div class="stat-label">Imágenes eliminadas</div></div><div class="stat-item"><div class="stat-value">${hasMultipleFiles?stats.totalParagraphs:stats.paragraphsCleaned}</div><div class="stat-label">Párrafos limpiados</div></div><div class="stat-item"><div class="stat-value">${stats.signatureRemoved?'Sí':'N/A'}</div><div class="stat-label">Firmas eliminadas</div></div>`;statsContainer.classList.remove("hidden")}function hideStats(){statsContainer.classList.add("hidden")}function resetForm(){selectedFiles=[];fileInput.value="";filePreview.classList.add("hidden");processBtn.disabled=false}
</script>
</body>
</html>'''

@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route("/health")
def health_check():
    return jsonify({"status": "healthy", "service": "CleanDoc", "version": "2.0.0"}), 200

@app.route("/limpiar_cedula", methods=["POST"])
def limpiar_endpoint():
    try:
        files = request.files.getlist("archivo")
        if not files or all(not f.filename for f in files):
            return jsonify({"error": "No se proporcionaron archivos"}), 400

        cleaned_files, stats_list = process_files(files)

        if not cleaned_files:
            return jsonify({"error": "No se pudieron procesar archivos válidos"}), 400

        if len(cleaned_files) == 1:
            return send_single_file(cleaned_files[0], stats_list[0])
        else:
            return send_multiple_files(cleaned_files, stats_list)

    except Exception as e:
        app.logger.error(f"Error: {e}")
        return jsonify({"error": "Error interno del servidor"}), 500

def process_files(files: List[FileStorage]) -> Tuple[List[Tuple[str, BytesIO]], List[CleaningStats]]:
    cleaner = DocumentCleaner()
    cleaned_files = []
    stats_list = []

    for file in files:
        if not file or not file.filename:
            continue

        try:
            safe_filename = secure_filename(file.filename)
            if not safe_filename.lower().endswith('.docx'):
                continue

            file.stream.seek(0)
            header = file.stream.read(4)
            file.stream.seek(0)
            if not header.startswith(b'PK\x03\x04'):
                continue

            cleaned_stream, stats = cleaner.clean_document(file.stream, safe_filename)
            cleaned_files.append((safe_filename, cleaned_stream))
            stats_list.append(stats)
        except Exception as e:
            app.logger.error(f"Error processing {file.filename}: {e}")
            continue

    return cleaned_files, stats_list

def send_single_file(file_data: Tuple[str, BytesIO], stats: CleaningStats):
    filename, stream = file_data
    response = send_file(stream, as_attachment=True, download_name=f"limpia_{filename}",
                        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response.headers['X-CleanDoc-Images-Removed'] = str(stats.images_removed)
    response.headers['X-CleanDoc-Paragraphs-Cleaned'] = str(stats.institutional_paragraphs_cleaned)
    response.headers['X-CleanDoc-Signature-Removed'] = str(stats.signature_section_removed)
    return response

def send_multiple_files(files_data: List[Tuple[str, BytesIO]], stats_list: List[CleaningStats]):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    try:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zf:
            for (filename, stream), stats in zip(files_data, stats_list):
                zf.writestr(f"limpia_{filename}", stream.read())
                stats_content = f"""CleanDoc - Estadísticas
Archivo: {filename}
Imágenes eliminadas: {stats.images_removed}
Párrafos limpiados: {stats.institutional_paragraphs_cleaned}
Textboxes limpiados: {stats.textboxes_cleaned}
Firmas eliminadas: {'Sí' if stats.signature_section_removed else 'No'}
Párrafos totales eliminados: {stats.paragraphs_removed}"""
                zf.writestr(f"limpia_{filename}_stats.txt", stats_content)

        tmp.seek(0)
        response = send_file(tmp.name, as_attachment=True, download_name="cleandoc_limpios.zip", mimetype="application/zip")
        total_images = sum(s.images_removed for s in stats_list)
        total_paragraphs = sum(s.institutional_paragraphs_cleaned for s in stats_list)
        response.headers['X-CleanDoc-Total-Files'] = str(len(files_data))
        response.headers['X-CleanDoc-Total-Images-Removed'] = str(total_images)
        response.headers['X-CleanDoc-Total-Paragraphs-Cleaned'] = str(total_paragraphs)
        return response
    except Exception as e:
        app.logger.error(f"Error creating ZIP: {e}")
        return jsonify({"error": "Error creando archivo ZIP"}), 500

@app.after_request
def add_security_headers(response):
    response.headers['X-Frame-Options'] = 'SAMEORIGIN'
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    return response

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"error": "Archivo demasiado grande", "message": "Máximo 50 MB"}), 413

@app.errorhandler(HTTPException)
def handle_http_exception(error):
    return jsonify({"error": error.name, "message": error.description}), error.code

@app.errorhandler(Exception)
def handle_unexpected_error(error):
    app.logger.error(f"Error inesperado: {error}", exc_info=True)
    return jsonify({"error": "Error interno del servidor"}), 500

if __name__ == "__main__":
    host = os.getenv('HOST', '0.0.0.0')
    port = int(os.getenv('PORT', '5001'))
    debug = os.getenv('DEBUG', 'False').lower() == 'true'
    app.run(host=host, port=port, debug=debug)
